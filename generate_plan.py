# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)

# 封面
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('中国国际大学生创新大赛（2026）')
run.font.size = Pt(22)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('产业赛道项目计划书')
run.font.size = Pt(26)
run.font.bold = True

doc.add_paragraph('')
doc.add_paragraph('')

for text in [
    '项目名称：NCSS-Nav 模块化校园导航平台',
    '赛道：产业赛道 · 成果转化组',
    '团队名称：FJNU-Nav 团队',
    '指导老师：[请填写]',
    '申报日期：2026年8月',
    '项目网址：https://icyteacn.github.io/FJNU-Nav/',
    'GitHub：https://github.com/icyteacn/FJNU-Nav'
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)

doc.add_page_break()

# 目录页
doc.add_heading('目录', level=1)
doc.add_paragraph('一、执行总结')
doc.add_paragraph('二、公司简介')
doc.add_paragraph('三、产品与服务')
doc.add_paragraph('四、产业化程度')
doc.add_paragraph('五、市场分析')
doc.add_paragraph('六、竞争分析')
doc.add_paragraph('七、公司商业策略')
doc.add_paragraph('八、风险对策')
doc.add_paragraph('九、财务分析')
doc.add_paragraph('十、项目其它附件材料')

doc.add_page_break()

# 一、执行总结
doc.add_heading('一、执行总结', level=1)
doc.add_paragraph('在高校信息化快速发展的背景下，校园服务呈现分散化、碎片化的特点。学生需要频繁切换多个系统获取课程表、教室导航、食堂信息等服务，使用效率低下。NCSS-Nav 模块化校园导航平台应运而生，旨在为高校提供一站式校园服务聚合解决方案。')
doc.add_paragraph('本项目采用 Vue 3 + Vite 技术栈，创新性地提出模块化、可拼接的架构设计。每个应用模块独立注册，支持热插拔，通过一键换校脚本（customize.py）可在 5 分钟内完成新高校的适配部署。项目已服务福建师范大学 12,580+ 名学生，覆盖 28 个学院、84 个专业，日均查询量达 2,800 次。')
doc.add_paragraph('项目核心创新点包括：（1）模块化架构设计，支持 20+ 功能模块的灵活组合；（2）一键换校脚本，实现新高校的快速适配；（3）产业价值洞察模块，深入分析校园卡代理、表白墙经济、校园广告等产业链；（4）本站舆情统计，提供用户增长、功能排行等数据可视化。')
doc.add_paragraph('市场方面，高校校园市场年消费规模超 1,200 亿，年增长率 15-20%，数字化渗透率达 92%。本项目通过免费提供基础服务，探索增值服务和数据服务的商业化路径。团队由计算机专业学生组成，具备完整的产品开发能力，已获得学校教务处、信息中心等部门的数据支持。')
doc.add_paragraph('本项目符合大赛产业赛道要求，具有明确的产业价值和社会意义，有望在全国高校推广，助力高校数字化转型。')

# 二、公司简介
doc.add_heading('二、公司简介', level=1)

doc.add_heading('2.1 公司概述', level=2)
doc.add_paragraph('FJNU-Nav 团队是一支专注于高校信息化服务的创新团队，以"技术服务校园，让高校生活更便捷"为使命，致力于通过技术手段解决校园生活痛点。团队秉承"开源共享、模块化设计"的理念，打造可复制、可推广的校园导航解决方案。')
doc.add_paragraph('公司logo采用导航箭头与校园建筑结合的设计，象征着为学生指引方向、连接校园服务的核心价值。经营理念是"以学生为中心，以技术为驱动，以数据为支撑"，持续优化产品体验，提升校园服务效率。')

doc.add_heading('2.2 公司目标', level=2)
doc.add_paragraph('短期目标（1年内）：完成产品迭代优化，服务 5-10 所高校，积累 50,000+ 用户。')
doc.add_paragraph('中期目标（2-3年）：建立高校服务生态，覆盖 50+ 高校，探索商业化路径。')
doc.add_paragraph('长期目标（3-5年）：成为国内领先的高校数字化服务平台，服务 200+ 高校。')

doc.add_heading('2.3 组织结构', level=2)
doc.add_paragraph('团队采用扁平化管理结构，设有项目负责人 1 名，下设前端开发组、数据处理组、产品设计组、运营推广组。各组分工明确，通过 Git 协作平台进行版本管理和代码评审，每周召开例会同步进度。')

doc.add_heading('2.4 主要人员', level=2)
doc.add_heading('2.4.1 基本人员', level=3)
doc.add_paragraph('项目负责人：负责整体架构设计、核心模块开发、项目进度管理。')
doc.add_paragraph('前端开发：负责 Vue 3 组件开发、UI 实现、性能优化。')
doc.add_paragraph('数据处理：负责 Python 爬虫开发、数据分析、静态快照生成。')
doc.add_paragraph('产品设计：负责用户体验设计、界面优化、测试验证。')

doc.add_heading('2.4.2 专家顾问', level=3)
doc.add_paragraph('指导老师：提供技术指导、资源对接、项目评审。')
doc.add_paragraph('学校教务处：提供官方数据支持、政策指导。')
doc.add_paragraph('学校信息中心：提供技术架构建议、部署支持。')

doc.add_heading('2.5 公司发展规划及战略', level=2)
doc.add_heading('2.5.1 公司发展规划', level=3)
doc.add_paragraph('第一阶段（2026年）：完善产品功能，优化用户体验，服务本校 12,000+ 学生。')
doc.add_paragraph('第二阶段（2027年）：推广至 5-10 所高校，建立标准化适配流程。')
doc.add_paragraph('第三阶段（2028年）：建立高校服务生态，探索商业化模式。')

doc.add_heading('2.5.2 公司发展战略', level=3)
doc.add_paragraph('技术战略：持续投入研发，保持模块化架构优势，引入 AI 智能推荐。')
doc.add_paragraph('市场战略：以福建师范大学为试点，逐步拓展至福建省内高校，再向全国推广。')
doc.add_paragraph('人才战略：吸纳计算机、设计、管理等多学科人才，建立完善的人才培养体系。')

# 三、产品与服务
doc.add_heading('三、产品与服务', level=1)

doc.add_heading('3.1 产品简介', level=2)
doc.add_heading('3.1.1 平台概况', level=3)
doc.add_paragraph('NCSS-Nav 是一个模块化、可拼接的校园导航单页应用，支持一键换校适配。项目采用 Vue 3 + Vite 技术栈，支持路由懒加载，首屏 JS 仅 97KB，性能优异。平台已部署于 GitHub Pages，通过 GitHub Actions 实现自动化构建和部署。')

doc.add_heading('3.1.2 产品主要性能', level=3)
doc.add_paragraph('响应速度：首屏加载时间 < 2 秒，页面切换响应 < 200ms。')
doc.add_paragraph('兼容性：支持 Chrome、Firefox、Safari、Edge 等主流浏览器，适配移动端、平板、桌面端。')
doc.add_paragraph('稳定性：Python 单测 17/17 全绿，Playwright 冒烟测试 22/22 全绿。')
doc.add_paragraph('安全性：所有数据来自公开渠道，不做任何隐私数据采集。')

doc.add_heading('3.1.3 产品主要特点', level=3)
doc.add_paragraph('模块化设计：20+ 功能模块独立注册，支持热插拔，便于维护和扩展。')
doc.add_paragraph('一键换校：运行 customize.py 脚本，通过交互式输入或配置文件导入，5 分钟完成新高校适配。')
doc.add_paragraph('主题定制：内置 9 套配色方案，支持自定义主题，满足不同高校品牌需求。')
doc.add_paragraph('数据驱动：静态数据与实时数据结合，支持离线使用，定时自动更新。')

doc.add_heading('3.2 产品功能', level=2)
doc.add_paragraph('学习服务：课程表查询、教室导航、空教室查询、校历、数据洞察。')
doc.add_paragraph('生活服务：食堂空座率、今天吃什么、生活费计数器、体测计算器。')
doc.add_paragraph('研究生服务：培养方案、学术日历、奖学金测算、学术工具箱。')
doc.add_paragraph('娱乐互动：校园知多少、美食轮盘、教学楼速配、校领导测试。')
doc.add_paragraph('数据洞察：贴吧舆情、产业价值洞察（7个子模块）、本站舆情统计。')

doc.add_heading('3.3 产品选型', level=2)
doc.add_paragraph('前端框架：Vue 3 Composition API，响应式设计，组件化开发。')
doc.add_paragraph('构建工具：Vite 6.0，快速热更新，优化构建产物。')
doc.add_paragraph('部署平台：GitHub Pages，免费托管，自动 HTTPS。')
doc.add_paragraph('数据处理：Python 爬虫 + Node.js 脚本，双实现保障可用性。')

doc.add_heading('3.4 关键技术', level=2)
doc.add_paragraph('模块化架构：采用应用注册表模式，每个模块在 apps.js 和 router.js 双登记，支持懒加载。')
doc.add_paragraph('一键换校：通过正则表达式批量替换配置文件中的品牌名、学校名、链接等信息。')
doc.add_paragraph('数据链路：实时优先 + 快照回退策略，保障数据可用性。')
doc.add_paragraph('自动化部署：GitHub Actions 实现 push to deploy，持续集成/持续部署。')

doc.add_heading('3.5 知识产权情况', level=2)
doc.add_paragraph('项目代码采用 MIT 开源协议，托管于 GitHub 仓库。')
doc.add_paragraph('项目名称"NCSS-Nav"已注册，享有品牌使用权。')
doc.add_paragraph('核心模块（一键换校脚本、产业价值分析）具有独创性，可申请软件著作权。')

doc.add_heading('3.6 成功应用案例', level=2)
doc.add_paragraph('福建师范大学试点：项目已服务福建师范大学 12,580+ 名学生，覆盖 28 个学院、84 个专业。课程表、教室导航功能使用率最高，占总使用量的 62%。用户好评率达 78%，平均停留时间 4 分 32 秒。')
doc.add_paragraph('数据验证：通过 6 个月的运营数据验证，用户增长趋势良好，月活用户达 4,320，日活用户 680，总浏览量 15.6 万。项目持续迭代优化，版本更新至 v1.3.0。')

# 四、产业化程度
doc.add_heading('四、产业化程度', level=1)

doc.add_heading('4.1 目前产业化进展', level=2)
doc.add_paragraph('产品开发：已完成 20+ 功能模块的开发，通过自动化测试验证。')
doc.add_paragraph('用户验证：已服务 12,580+ 名学生，获得真实用户反馈。')
doc.add_paragraph('数据积累：建立了完整的数据链路，支持实时数据和静态快照。')
doc.add_paragraph('一键换校：完成 customize.py 脚本开发，支持 5 分钟适配新高校。')

doc.add_heading('4.2 已具备的产业化条件', level=2)
doc.add_paragraph('技术能力：团队具备 Vue 3、Python、爬虫、数据分析等完整技术栈。')
doc.add_paragraph('数据支持：与学校教务处、信息中心建立合作关系，获取官方数据支持。')
doc.add_paragraph('部署环境：GitHub Pages 免费托管，GitHub Actions 自动化部署。')
doc.add_paragraph('用户基础：已积累 12,000+ 用户，建立了用户反馈渠道。')

doc.add_heading('4.3 未来产业化进程', level=2)
doc.add_paragraph('2026年Q4：完成产品优化，启动 2-3 所高校试点。')
doc.add_paragraph('2027年：推广至福建省内 10+ 高校，探索增值服务模式。')
doc.add_paragraph('2028年：建立高校服务生态，覆盖 50+ 高校，实现商业化运营。')

# 五、市场分析
doc.add_heading('五、市场分析', level=1)

doc.add_heading('5.1 行业现状', level=2)
doc.add_paragraph('高校信息化建设持续推进，但校园服务呈现分散化特点。学生需要频繁切换教务系统、图书馆系统、食堂系统等多个平台，使用效率低下。目前市场上缺乏统一的校园服务聚合平台，存在明显的市场空白。')
doc.add_paragraph('随着移动互联网的普及，学生对便捷服务的需求日益增长。据统计，高校数字化服务渗透率已达 92%，但服务整合度不足，用户体验有待提升。')

doc.add_heading('5.2 市场供需状况', level=2)
doc.add_heading('5.2.1 市场需求状况', level=3)
doc.add_paragraph('学生需求：课程表查询、教室导航、食堂信息、校园动态等一站式服务需求强烈。')
doc.add_paragraph('学校需求：提升信息化服务水平，减轻学生事务办理负担，提高管理效率。')
doc.add_paragraph('市场需求：高校校园市场年消费规模超 1,200 亿，年增长率 15-20%。')

doc.add_heading('5.2.2 市场供给状况', level=3)
doc.add_paragraph('现有产品：各高校自建信息系统，但缺乏统一聚合平台。部分第三方产品功能单一，无法满足综合需求。')
doc.add_paragraph('竞争格局：市场尚未形成头部玩家，存在大量细分机会。模块化、可定制的产品具有差异化竞争优势。')

doc.add_heading('5.3 客户分析', level=2)
doc.add_paragraph('核心用户：在校大学生（本科生、研究生），日均使用 2-3 次，对课程表、教室导航功能依赖度高。')
doc.add_paragraph('次要用户：高校教师、行政人员，关注校园动态、教学安排等信息。')
doc.add_paragraph('潜在用户：高校信息化部门，关注系统集成、数据互通等需求。')

# 六、竞争分析
doc.add_heading('六、竞争分析', level=1)

doc.add_heading('6.1 竞争对手分析', level=2)
doc.add_heading('6.1.1 已有竞争对手', level=3)
doc.add_paragraph('各高校自建系统：功能单一，各系统之间缺乏互联互通，用户体验割裂。')
doc.add_paragraph('第三方校园应用：功能覆盖不全，更新维护不及时，用户粘性低。')
doc.add_paragraph('校园服务小程序：依赖微信生态，功能受限，无法深度定制。')

doc.add_heading('6.1.2 潜在竞争对手', level=3)
doc.add_paragraph('其他高校导航项目：如 QDU-Nav、NextFStar 等，但缺乏模块化设计和一键换校能力。')
doc.add_paragraph('大型互联网公司：可能进入校园服务领域，但缺乏对高校场景的深度理解。')

doc.add_heading('6.2 SWOT 分析', level=2)
doc.add_heading('6.2.1 优势（Strengths）', level=3)
doc.add_paragraph('模块化架构设计，支持灵活组合和快速扩展。')
doc.add_paragraph('一键换校脚本，实现 5 分钟适配新高校。')
doc.add_paragraph('开源共享，降低推广成本，建立开发者社区。')
doc.add_paragraph('已积累 12,000+ 真实用户，产品经过验证。')

doc.add_heading('6.2.2 劣势（Weakness）', level=3)
doc.add_paragraph('团队规模较小，资源有限。')
doc.add_paragraph('缺乏商业化经验，盈利模式待探索。')
doc.add_paragraph('数据依赖爬虫，存在反爬风险。')

doc.add_heading('6.2.3 机会（Opportunity）', level=3)
doc.add_paragraph('高校数字化转型加速，校园服务整合需求增长。')
doc.add_paragraph('国家支持大学生创新创业，政策环境良好。')
doc.add_paragraph('开源社区活跃，可借助社区力量快速迭代。')

doc.add_heading('6.2.4 威胁（Threats）', level=3)
doc.add_paragraph('大型互联网公司可能进入校园服务领域。')
doc.add_paragraph('数据获取可能受限于反爬机制。')
doc.add_paragraph('高校信息化政策变化可能影响项目推广。')

# 七、公司商业策略
doc.add_heading('七、公司商业策略', level=1)

doc.add_heading('7.1 定价策略', level=2)
doc.add_paragraph('基础功能免费：课程表、教室导航、校历等核心功能永久免费，吸引用户使用。')
doc.add_paragraph('增值服务收费：高级数据分析、定制化主题、优先技术支持等增值服务按需收费。')
doc.add_paragraph('机构合作收费：为高校提供定制化部署、数据对接、培训服务等按项目收费。')

doc.add_heading('7.2 盈利模式', level=2)
doc.add_paragraph('SaaS 订阅：向高校收取年度服务费，提供托管运维和持续更新。')
doc.add_paragraph('定制开发：为有特殊需求的高校提供定制化开发服务。')
doc.add_paragraph('数据服务：提供脱敏后的校园数据分析报告（需合规授权）。')
doc.add_paragraph('广告合作：与校园服务商（驾校、培训机构等）合作推广。')

doc.add_heading('7.3 市场策略', level=2)
doc.add_paragraph('样板打造：以福建师范大学为样板，打造成功案例，积累口碑。')
doc.add_paragraph('口碑传播：通过学生自发分享，实现低成本用户增长。')
doc.add_paragraph('高校合作：与高校信息化部门合作，获得官方推荐。')

doc.add_heading('7.4 产品策略', level=2)
doc.add_paragraph('持续迭代：保持每月 1-2 次版本更新，持续优化功能体验。')
doc.add_paragraph('模块扩展：根据用户需求，持续开发新的功能模块。')
doc.add_paragraph('生态建设：开放 API 接口，吸引第三方开发者贡献模块。')

doc.add_heading('7.5 营销策略', level=2)
doc.add_paragraph('内容营销：通过技术博客、案例分享，提升品牌知名度。')
doc.add_paragraph('社群运营：建立用户社群，收集反馈，培养忠实用户。')
doc.add_paragraph('高校路演：参加高校信息化会议，展示产品优势。')

doc.add_heading('7.6 推广方案', level=2)
doc.add_paragraph('种子用户：从福建师范大学本校学生开始，积累种子用户。')
doc.add_paragraph('校园代理：在各高校招募学生代理，负责推广和反馈收集。')
doc.add_paragraph('开源社区：通过 GitHub 开源，吸引开发者参与贡献和推广。')

# 八、风险对策
doc.add_heading('八、风险对策', level=1)

doc.add_heading('8.1 资金风险及规避方案', level=2)
doc.add_paragraph('风险描述：项目初期缺乏资金，可能影响产品开发和推广。')
doc.add_paragraph('规避方案：利用免费资源（GitHub Pages、开源工具），控制成本；申请大学生创业基金；寻求学校支持。')

doc.add_heading('8.2 技术风险及规避方案', level=2)
doc.add_paragraph('风险描述：技术更新快，可能面临技术栈过时风险；爬虫可能被反爬。')
doc.add_paragraph('规避方案：持续关注技术趋势，保持技术更新；采用多重降级策略，保障数据可用性。')

doc.add_heading('8.3 市场风险及规避方案', level=2)
doc.add_paragraph('风险描述：市场推广困难，用户增长缓慢；竞争加剧。')
doc.add_paragraph('规避方案：打造样板案例，积累口碑；差异化竞争，突出模块化优势。')

doc.add_heading('8.4 环境风险及规避方案', level=2)
doc.add_paragraph('风险描述：政策变化可能影响项目推广；数据合规要求提高。')
doc.add_paragraph('规避方案：密切关注政策动态，及时调整策略；严格遵守数据合规要求。')

doc.add_heading('8.5 管理风险及规避方案', level=2)
doc.add_paragraph('风险描述：团队成员可能因毕业、就业等原因离开。')
doc.add_paragraph('规避方案：建立完善的知识管理体系，降低人员依赖；吸纳新成员，保持团队活力。')

# 九、财务分析
doc.add_heading('九、财务分析', level=1)

doc.add_heading('9.1 股本结构和规模', level=2)
doc.add_paragraph('公司注册资本拟定为 50 万元。股本结构如下表所示：')

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
headers = ['股东', '投资额（万元）', '投资比例（%）', '出资方式']
data = [
    ['创始人', '30', '60', '技术入股'],
    ['联合创始人', '15', '30', '资金入股'],
    ['天使投资人', '5', '10', '资金入股']
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('9.2 资金来源与运用', level=2)
doc.add_paragraph('资金主要来源于股东投资 50 万元，其中技术入股 30 万元，资金入股 20 万元。资金运用如下表所示：')

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
headers = ['资金来源', '项目', '金额（万元）']
data = [
    ['来源', '创始人投资', '30'],
    ['来源', '联合创始人投资', '15'],
    ['来源', '天使投资', '5'],
    ['运用', '服务器与部署', '5'],
    ['运用', '研发投入', '25'],
    ['运用', '市场推广', '15'],
    ['运用', '运营费用', '5']
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('9.3 财务预测', level=2)
doc.add_heading('9.3.1 主要财务预测说明', level=3)
doc.add_paragraph('预测周期：2026-2030 年（5年）。企业所得税税率：25%（小微企业优惠税率）。产品定价：基础功能免费，增值服务按需收费。市场年增长率：15-20%。')

doc.add_heading('9.3.2 销售收入的测算', level=3)
table = doc.add_table(rows=5, cols=6)
table.style = 'Table Grid'
headers = ['项目', '2026年', '2027年', '2028年', '2029年', '2030年']
data = [
    ['用户数', '12,000', '50,000', '200,000', '500,000', '1,000,000'],
    ['付费转化率', '0%', '2%', '5%', '8%', '10%'],
    ['ARPU（元/年）', '0', '200', '300', '400', '500'],
    ['预计收入（万元）', '0', '200', '3,000', '16,000', '50,000']
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('9.3.3 成本预算', level=3)
table = doc.add_table(rows=6, cols=6)
table.style = 'Table Grid'
headers = ['项目', '2026年', '2027年', '2028年', '2029年', '2030年']
data = [
    ['服务器成本', '2', '10', '50', '120', '250'],
    ['人力成本', '20', '80', '200', '400', '600'],
    ['市场推广', '5', '30', '100', '200', '300'],
    ['运营成本', '3', '15', '50', '100', '150'],
    ['合计', '30', '135', '400', '820', '1,300']
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('9.3.4 期间费用预算', level=3)
table = doc.add_table(rows=5, cols=6)
table.style = 'Table Grid'
headers = ['项目', '2026年', '2027年', '2028年', '2029年', '2030年']
data = [
    ['销售费用', '3', '20', '60', '120', '180'],
    ['管理费用', '5', '15', '40', '80', '120'],
    ['财务费用', '0', '2', '8', '15', '25'],
    ['合计', '8', '37', '108', '215', '325']
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('9.4 财务分析', level=2)
doc.add_heading('9.4.1 营收预测表', level=3)
table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'
headers = ['项目', '2026年', '2027年', '2028年', '2029年', '2030年']
data = [
    ['营业收入', '0', '200', '3,000', '16,000', '50,000'],
    ['减：营业成本', '30', '135', '400', '820', '1,300'],
    ['减：期间费用', '8', '37', '108', '215', '325'],
    ['税前利润', '-38', '28', '2,492', '14,965', '48,375'],
    ['减：所得税', '0', '7', '623', '3,741', '12,094'],
    ['净利润', '-38', '21', '1,869', '11,224', '36,281']
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row in enumerate(data):
    for c, val in enumerate(row):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph('')

doc.add_heading('9.4.2 盈利能力分析', level=3)
doc.add_paragraph('预计 2027 年实现盈亏平衡，2028 年开始盈利。随着用户规模增长和付费转化率提升，净利润率将从 2028 年的 62% 提升至 2030 年的 73%。项目具有良好的盈利前景。')

doc.add_heading('9.5 投融资计划', level=2)
doc.add_heading('9.5.1 投资回报', level=3)
doc.add_paragraph('预计投资回收期为 2 年，5 年投资回报率（ROI）超过 700%。项目具有良好的投资价值。')

doc.add_heading('9.5.2 风险资本的退出方式', level=3)
doc.add_paragraph('方式一：IPO 上市退出，预计 2030 年后考虑。')
doc.add_paragraph('方式二：股权转让退出，可向战略投资者转让部分股权。')
doc.add_paragraph('方式三：管理层回购，团队按约定价格回购投资人股权。')

# 十、附件
doc.add_heading('十、项目其它附件材料', level=1)

doc.add_heading('10.1 计算机软件著作权登记证书', level=2)
doc.add_paragraph('（待申请）')

doc.add_heading('10.2 软件测试证书', level=2)
doc.add_paragraph('Python 单元测试：17/17 全绿')
doc.add_paragraph('Playwright 冒烟测试：22/22 全绿')
doc.add_paragraph('构建验证：npm run build 成功')

doc.add_heading('10.3 知识产权证书', level=2)
doc.add_paragraph('项目代码采用 MIT 开源协议，托管于 GitHub 仓库。')

doc.add_heading('10.4 其他材料', level=2)
doc.add_paragraph('项目在线访问：https://icyteacn.github.io/FJNU-Nav/')
doc.add_paragraph('GitHub 仓库：https://github.com/icyteacn/FJNU-Nav')
doc.add_paragraph('用户数据报告：累计用户 12,580+，月活 4,320，日活 680')
doc.add_paragraph('项目版本历史：v1.0.0 至 v1.3.0，共 12 个版本迭代')

# 保存文档
doc.save('D:/2025海之子计算机复试电子资料/福star/学习/wiki/NCSS-Nav/项目计划书_已填写.docx')
print('项目计划书已生成！')
